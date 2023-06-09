import tkinter as tk
from tkinter import ttk
from tkinter import *
from tkinter import filedialog
from tkinter import messagebox
from docx import Document
from docx.shared import Pt
import win32com.client as win32
import json
import os
from os.path import exists
from fpdf import FPDF

def tabla(datos):
    ventana = tk.Tk()
    ventana.title("Tabla de datos")

    # Crear una tabla utilizando Treeview de ttk
    tabla = ttk.Treeview(ventana, selectmode='browse')
    tabla['columns'] = tuple(datos.columns)
    tabla['show'] = 'headings'

    # Configurar encabezados de columna
    for columna in datos.columns:
        tabla.heading(columna, text=columna)
        tabla.column(columna, width=100)

    # Agregar filas a la tabla
    for fila in datos.itertuples(index=False):
        tabla.insert('', 'end', values=tuple(fila))

    # Mostrar la tabla en la ventana
    tabla.pack(expand=True, fill=tk.BOTH)

    # # Boton para guardar en excel
    bottomFrame = Frame(ventana)
    boton_guardarExcel = Button(bottomFrame, text="Excel")
    boton_guardarExcel.bind("<Button>", (lambda event: guardarExcel(datos)))
    boton_guardarExcel.pack(side=LEFT)
    boton_guardarWord = Button(bottomFrame, text='word')
    boton_guardarWord.bind("<Button>", (lambda event: guardarWord(datos)))
    boton_guardarWord.pack(side=LEFT)
    boton_guardarCSV = Button(bottomFrame, text='CSV')
    boton_guardarCSV.bind("<Button>", (lambda event: guardarCSV(tabla, datos)))
    boton_guardarCSV.pack(side=LEFT)
    boton_guardarJSON = Button(bottomFrame, text='JSON')
    boton_guardarJSON.bind("<Button>", (lambda event: serializarTablaJson(datos)))
    boton_guardarJSON.pack(side=LEFT)
    boton_GuardarPDF = Button(bottomFrame, text='PDF')
    boton_GuardarPDF.bind("<Button>", (lambda event: guardarPDF(datos)))
    boton_GuardarPDF.pack(side=LEFT)

    bottomFrame.pack(side=TOP)
    ventana.mainloop()

def guardarExcel(datos):
    archivo = filedialog.asksaveasfilename(defaultextension=".xlsx",
                                           filetypes=[("Archivo de Excel", "*.xlsx"), ("Todos los archivos", "*.*")])
    if not archivo:
        return
    if exists(archivo):
        messagebox.showinfo(message='El archivo ya existe, por favor vuelva a intentar', title='error')
        return
    datos.to_excel(archivo, index=False)
    os.system(archivo)    

def guardarWord(df):
    word_app = win32.gencache.EnsureDispatch('Word.Application')
    word_app.Visible = True

    doc = word_app.Documents.Add()
    doc.Content.Font.Name = 'Arial'
    doc.Content.Font.Size = 5

    tabla = doc.Tables.Add(doc.Range(0, 0), len(df) + 1, len(df.columns))
    tabla.Style = 'Table Grid'  

    encabezados = df.columns
    for i, encabezado in enumerate(encabezados):
        tabla.Cell(1, i + 1).Range.Text = encabezado

    for i, fila in df.iterrows():
        for j, valor in enumerate(fila):
            tabla.Cell(i + 2, j + 1).Range.Text = str(valor)

    archivo = filedialog.asksaveasfilename(defaultextension=".docx",
                                        filetypes=[("Archivo de Word", "*.docx"), ("Todos los archivos", "*.*")])
    if not archivo:
        return
    if exists(archivo):
        messagebox.showinfo(message='El archivo ya existe, por favor vuelva a intentar', title='error')
        return
    
    doc.SaveAs(archivo)
    word_app.Quit()

def guardarPDF(df):
    
# Crear un nuevo documento PDF
    doc = FPDF()

    # Agregar una página al documento
    doc.add_page()

    # Configurar la fuente y el tamaño para el contenido de la tabla
    doc.set_font('Arial', '', 6)

    # Agregar encabezados de columna a la tabla
    for col in df.columns:
        doc.cell(30, 10, col, 1, 0, 'C')
    doc.ln()

    # Agregar filas al DataFrame
    for _, fila in df.iterrows():
        for valor in fila:
            doc.cell(30, 10, str(valor), 1, 0, 'C')
        doc.ln()
    archivo = filedialog.asksaveasfilename(defaultextension=".pdf",
                                            filetypes=[("Archivo de PDF", "*.pdf"), ("Todos los archivos", "*.*")])
    if not archivo:
        return
    if exists(archivo):
        messagebox.showinfo(message='El archivo ya existe, por favor vuelva a intentar', title='error')
        return
    # Guardar el documento PDF
    doc.output(archivo)

# def guardarWord(df):
#     archivo = filedialog.asksaveasfilename(defaultextension=".docx",
#                                            filetypes=[("Archivo de Word", "*.docx"), ("Todos los archivos", "*.*")])
#     doc = Document()

#     # Agregar una tabla al documento
#     tabla = doc.add_table(rows=1, cols=len(df.columns))
#     tabla.style = 'Table Grid'  # Aplicar estilo de tabla

#     # Establecer el estilo de fuente para toda la tabla
#     fuente = 'Arial'
#     tamaño_fuente = Pt(5)

#     # Agregar encabezados de columna a la tabla
#     encabezados = df.columns
#     for i, encabezado in enumerate(encabezados):
#         celda = tabla.cell(0, i)
#         párrafo = celda.paragraphs[0]
#         párrafo.add_run(encabezado)

#     # Agregar filas al DataFrame
#     for _, fila in df.iterrows():
#         nueva_fila = tabla.add_row().cells
#         for i, valor in enumerate(fila):
#             párrafo = nueva_fila[i].paragraphs[0]
#             párrafo.add_run(str(valor))

#     for celda in tabla.rows[0].cells:
#         for párrafo in celda.paragraphs:
#             for run in párrafo.runs:
#                 run.font.name = fuente
#                 run.font.size = tamaño_fuente

#     # Guardar el documento
#     doc.save(archivo)

def guardarCSV(tabla, datos):
    # Actualizar los datos en el DataFrame de Pandas
    for fila in tabla.get_children():
        valores = tabla.item(fila)['values']
        indice = tabla.index(fila)
        for col, val in enumerate(valores):
            datos.iat[indice, col] = val
    # Guardar el DataFrame actualizado en un archivo
    archivo = filedialog.asksaveasfilename(defaultextension=".csv",
                                        filetypes=[("Archivo de CSV", "*.csv"), ("Todos los archivos", "*.*")])
    if not archivo:
        return
    if exists(archivo):
        messagebox.showinfo(message='El archivo ya existe, por favor vuelva a intentar', title='error')
        return
    datos.to_csv(archivo, index=False)
    os.system(archivo)

def serializarTablaJson(datos):
    # Convertir la tabla a un diccionario
    tabla_dict = {"columnas": list(datos.columns), "filas": datos.values.tolist()}

    # Serializar el diccionario a JSON
    tabla_json = json.dumps(tabla_dict, indent=4)

    archivo = filedialog.asksaveasfilename(defaultextension=".json",
                                        filetypes=[("Archivo de JSON", "*.json"), ("Todos los archivos", "*.*")])
    if not archivo:
        return
    if exists(archivo):
        messagebox.showinfo(message='El archivo ya existe, por favor vuelva a intentar', title='error')
        return
    
    # Guardar el JSON en un archivo
    with open(archivo, "w") as archivo:
        archivo.write(tabla_json)
    messagebox.showinfo(message='Archivo Guardado')